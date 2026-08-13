from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Token Plan 产品文档.docx"
OUT = ROOT / "Token Plan 产品文档（完整版）.docx"
FAQ_SOURCE = ROOT / "src/pages/resource-subscription/token-plan-faq.ts"
NAVY = RGBColor(23, 54, 93)
MUTED = RGBColor(92, 107, 128)


def style_run(run, size=10.5, bold=False, color=None, mono=False):
    name = "Consolas" if mono else "Microsoft YaHei"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), bold=True, color=NAVY)
        style_run(p.add_run(text[len(bold_prefix):]))
    else:
        style_run(p.add_run(text))
    return p


def item(doc, text, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(text))
    return p


def placeholder(doc, title, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run(f"【截图占位：{title}】"), 9.5, True, NAVY)
    style_run(p.add_run(description), 9.5, color=MUTED)


def code_block(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.right_indent = Pt(14)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(value.splitlines()):
        if idx:
            p.add_run().add_break()
        style_run(p.add_run(line), 8.5, color=RGBColor(36, 54, 75), mono=True)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Pt(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "2563EB"); cell._tc.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(header), 9.5, True, RGBColor(255, 255, 255))
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Pt(width)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri % 2:
                shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F7F9FC"); cells[i]._tc.get_or_add_tcPr().append(shd)
            style_run(cells[i].paragraphs[0].add_run(str(value)), 9.2, color=RGBColor(36, 54, 75))
    doc.add_paragraph()
    return table


def quick_start(doc, personal):
    edition = "个人版" if personal else "企业版"
    h(doc, f"{edition}快速开始", 1)
    para(doc, "完成订阅并获取接入信息后，即可发起首次模型调用。" if personal else "企业管理员完成购买、添加成员和分配席位后，成员即可获取接入信息并调用模型。")

    h(doc, f"步骤一：订阅 Token Plan {edition}", 2)
    item(doc, f"进入 Token Plan {edition}购买页面。", True)
    item(doc, "选择订阅周期和套餐档位，确认订单后完成支付。" if personal else "选择订阅周期和套餐档位，填写各档位席位数量，确认订单后完成支付。", True)
    item(doc, "支付成功后订阅立即生效。", True)

    if not personal:
        h(doc, "步骤二：添加成员并分配席位", 2)
        item(doc, "进入左侧导航的“成员管理”，添加需要使用 Token Plan 的企业成员；已有成员可跳过此操作。", True)
        item(doc, "进入“Token Plan > 企业订阅”，在席位管理中选择未分配席位，并将席位分配给对应成员。", True)
        item(doc, "每个席位同一时间只能分配给一名成员。分配完成后，该成员可使用席位对应的套餐额度。", True)
        placeholder(doc, "成员管理—添加成员", "展示成员管理页面的添加成员入口，以及成员添加完成后的列表位置。")
        placeholder(doc, "企业订阅—分配席位", "展示未分配席位、分配成员入口和确认分配操作。")

    step = "二" if personal else "三"
    h(doc, f"步骤{step}：获取 API Key 和 Base URL", 2)
    item(doc, ("购买者" if personal else "已获配席位的成员") + "进入“Token Plan > 我的订阅”。", True)
    item(doc, "在“配置”区域生成并复制专属 API Key。完整 Key 请妥善保存，不要公开或与他人共享。", True)
    item(doc, "复制页面展示的 Base URL，用于后续接口或客户端配置。", True)
    para(doc, "Base URL：https://neolink.com/api/v1")
    placeholder(doc, "我的订阅—配置", "展示专属 API Key 的生成或复制入口，以及 Base URL 所在位置。截图时请隐藏完整 API Key。")
    para(doc, "注意：订阅专用 API Key 仅用于 Token Plan 专属地址。重置 Key 后，旧 Key 会立即失效，请同步更新所有调用配置。")

    step = "三" if personal else "四"
    h(doc, f"步骤{step}：发起首次调用", 2)
    para(doc, "将以下示例中的 API Key 和模型名称替换为您自己的配置后发起请求。请求中的 model 必须是当前套餐支持的模型。")
    code_block(doc, '''curl https://neolink.com/api/v1/chat/completions \\
  -H "Authorization: Bearer <YOUR_TOKEN_PLAN_API_KEY>" \\
  -H "Content-Type: application/json" \\
  -d '{
    "model": "deepseek-v4-pro",
    "messages": [{"role": "user", "content": "你好"}]
  }' ''')
    para(doc, "调用成功后，可在“我的订阅”查看额度变化和使用明细。")
    para(doc, "相关文档：模型接口文档、客户端工具接入指南。")


def team_management(doc):
    h(doc, "企业版团队管理", 1)
    para(doc, "企业管理员可添加成员、分配和回收席位，并查看团队席位的使用情况。")
    h(doc, "成员管理", 2)
    h(doc, "添加成员", 3)
    for x in ["进入左侧导航的“成员管理”。", "添加需要使用 Token Plan 的企业成员；成员已存在时无需重复添加。", "添加完成后，进入“Token Plan > 企业订阅”为其分配席位。"]:
        item(doc, x, True)
    placeholder(doc, "成员管理—添加成员", "展示添加成员入口、成员信息填写区域和成员列表。")
    h(doc, "移除或禁用成员", 3)
    item(doc, "禁用成员：API Key 停止调用，但已分配的 Token Plan 席位仍保持占用；重新启用后可恢复使用。")
    item(doc, "移除成员：成员将离开企业，其已分配的 Token Plan 席位自动回收并恢复为未分配状态，可重新分配给其他成员。")

    h(doc, "席位管理", 2)
    para(doc, "进入“Token Plan > 企业订阅”查看全部席位。企业管理员可按席位状态筛选，并进行分配、回收、升级和查看详情等操作。")
    placeholder(doc, "企业订阅—席位列表", "展示席位档位、分配成员、使用状态、剩余额度和操作入口。")
    h(doc, "分配席位", 3)
    for x in ["找到状态为“未分配”的席位，点击“分配”。", "选择一名企业成员并确认。", "分配完成后，成员可进入“我的订阅”生成专属 API Key。"]:
        item(doc, x, True)
    para(doc, "注意：每个成员同一时间只能持有一个席位，每个席位同一时间只能分配给一名成员。")
    placeholder(doc, "分配席位", "展示成员选择、席位信息和确认分配操作。")
    h(doc, "回收和重新分配席位", 3)
    para(doc, "在席位列表中找到已分配席位，点击“回收”。回收后席位变为未分配状态，原成员不能再使用该席位，管理员可以将其重新分配给其他成员。")
    h(doc, "升级席位", 3)
    para(doc, "选择目标席位并点击“升级”，选择更高档位后提交订单。具体升级规则和费用以套餐概览及订单页面为准。")
    h(doc, "查看使用情况", 2)
    for x in ["在席位列表中查看每个席位的档位、分配成员、额度使用进度和到期时间。", "点击“查看用量明细”，查看该席位的用量记录。", "企业管理员可在“资源统计”中按成员、模型和时间范围分析团队用量。"]:
        item(doc, x)
    placeholder(doc, "席位使用情况", "展示额度使用进度、剩余额度和“查看用量明细”按钮。")


def read_faqs():
    text = FAQ_SOURCE.read_text(encoding="utf-8")
    pattern = re.compile(r'\{ id: "([^"]+)", group: "([^"]+)", question: "([^"]+)", answer: "([^"]+)".*?editions: ([^}\n]+)')
    faqs = []
    for m in pattern.finditer(text):
        editions = m.group(5)
        faqs.append({"id": m.group(1), "group": m.group(2), "q": m.group(3), "a": m.group(4), "personal": "BOTH" in editions or "personal" in editions, "enterprise": "BOTH" in editions or "enterprise" in editions})
    return faqs


def faq_page(doc, edition, faqs):
    h(doc, f"{edition}常见问题", 1)
    key = "personal" if edition == "个人版" else "enterprise"
    groups = [("purchase", "套餐购买与续费"), ("credit", "Credit 与用量"), ("access", "API Key 与接入"), ("team", "成员与席位")]
    for group_key, title in groups:
        items = [x for x in faqs if x[group_key if False else key] and x["group"] == group_key]
        if not items:
            continue
        h(doc, title, 2)
        for x in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            style_run(p.add_run(x["q"]), 10.5, True, NAVY)
            para(doc, x["a"])


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Pt(52)
section.left_margin = section.right_margin = Pt(56)
for name, size, color in (("Normal", 10.5, None), ("Title", 30, NAVY), ("Subtitle", 13, MUTED), ("Heading 1", 17, NAVY), ("Heading 2", 13, RGBColor(37, 99, 235)), ("Heading 3", 11, NAVY)):
    st = doc.styles[name]
    st.font.name = "Microsoft YaHei"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); st.font.size = Pt(size)
    if color: st.font.color.rgb = color
    if name.startswith("Heading"): st.font.bold = True
doc.styles["Normal"].paragraph_format.line_spacing = 1.22
doc.styles["Normal"].paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(85)
style_run(p.add_run("TOKEN PLAN"), 14, True, RGBColor(37, 99, 235))
p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("产品文档（完整版）")
p = doc.add_paragraph(style="Subtitle"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("页面文案审阅稿｜企业版 · 个人版")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(22); style_run(p.add_run("更新日期：2026 年 8 月"), 9.5, color=MUTED)
doc.add_page_break()

pricing = [("Lite 轻量版", "原价 ¥198/席/月\n活动价 ¥149/席/月", "40,000 Credits/席/月", "适合日常轻量使用与个人效率辅助"), ("Standard 标准版", "原价 ¥598/席/月\n活动价 ¥419/席/月", "120,000 Credits/席/月", "适合日常开发、办公与稳定使用"), ("Pro 尊享版", "原价 ¥1,398/席/月\n活动价 ¥909/席/月", "280,000 Credits/席/月", "适合高频调用与核心生产任务")]
models = [("DeepSeek", "deepseek-v4-pro", "推理模型、文本生成"), ("DeepSeek", "deepseek-v4-flash", "推理模型、文本生成"), ("Kimi", "kimi-k2.7-code", "推理模型、视觉理解、文本生成"), ("Kimi", "kimi-k2.6", "推理模型、视觉理解、文本生成"), ("Kimi", "kimi-k2.5", "推理模型、视觉理解、文本生成"), ("智谱", "glm-5.2", "推理模型、文本生成"), ("智谱", "glm-5.1", "推理模型、文本生成"), ("智谱", "glm-5", "推理模型、文本生成"), ("阿里", "qwen3.8-max-preview", "推理模型、视觉理解、文本生成"), ("阿里", "qwen3.7-max", "推理模型、文本生成"), ("阿里", "qwen3.7-plus", "推理模型、视觉理解、文本生成"), ("阿里", "qwen3.6-plus", "推理模型、视觉理解、文本生成"), ("阿里", "qwen3.6-flash", "推理模型、视觉理解、文本生成"), ("阿里", "qwen-image-2.0", "图片生成"), ("阿里", "qwen-image-2.0-pro", "图片生成"), ("阿里", "wan2.7-image", "图片生成"), ("阿里", "wan2.7-image-pro", "图片生成"), ("阿里", "happyhorse-1.1-i2v", "视频生成"), ("阿里", "happyhorse-1.1-t2v", "视频生成"), ("阿里", "happyhorse-1.1-r2v", "视频生成"), ("MiniMax", "MiniMax-M2.5", "推理模型、文本生成")]

def overview(doc, personal):
    edition = "个人版" if personal else "企业版"
    h(doc, f"Token Plan {edition}套餐概览", 1)
    para(doc, f"Token Plan {edition}面向" + ("个人开发者" if personal else "企业和团队") + "，提供 Lite 轻量版、Standard 标准版和 Pro 尊享版三个档位，匹配不同使用强度。")
    h(doc, "产品特点", 2)
    features = ["主流模型一站式使用：一个订阅即可使用套餐范围内的多种主流模型，按需切换，无需分别采购和配置。", "兼容主流 AI 工具：使用 Token Plan 专属 API Key，可接入支持自定义模型服务的 AI 编程与智能体工具。", "多档套餐灵活选择：提供 Lite 轻量版、Standard 标准版和 Pro 尊享版三个档位，覆盖从日常辅助到高频使用的不同需求。", "订阅用量清晰可控：支持按月或按年订阅，以 Credit 统一计量，可在控制台查看额度和使用明细。"]
    if not personal: features += ["按需组合团队席位：不同档位席位可按成员使用强度灵活搭配，兼顾使用需求与团队预算。", "团队订阅统一管理：管理员可统一分配、回收和升级席位，并查看成员用量；每个席位独享额度并对应一个专属 API Key。"]
    for x in features: item(doc, x)
    h(doc, "产品定价", 2)
    para(doc, "选择套餐类型、数量和订阅周期后完成购买。支持按月、按年购买，也可开启连续包月或连续包年。")
    if not personal:
        h(doc, "席位说明", 3)
        para(doc, "席位是 Token Plan 企业版的基础订阅单位，代表一名团队成员的使用名额。管理员将席位分配给成员后，成员可使用一个专属 API Key 调用服务。每个席位仅可绑定一名成员，不支持多人共用。")
    add_table(doc, ["套餐类型" if personal else "席位类型", "月付价格", "每月额度", "适用场景"], pricing, [92, 130, 135, 180])
    para(doc, ("升级" if personal else "加购或升级席位") + "时，费用按当前订阅的剩余周期折算，实际支付金额以订单页面展示为准。")
    h(doc, "Credit 抵扣说明", 2)
    para(doc, "模型调用统一折算为 Credit，并从当前订阅或席位的周期额度中扣减。单次消耗由模型类型、Token 用量、思考模式及工具调用等因素决定，实际消耗以控制台用量明细为准。")
    para(doc, "以 glm-5.2 为例，单次请求的 Credit 消耗示例如下（不同模型的单价不同，实际以账单为准）：")
    add_table(doc, ["Token 类型", "数量", "消耗 Credits"], [("输入 Tokens", "5,000", "8.00"), ("缓存 Tokens", "20,000", "8.00"), ("输出 Tokens", "1,000", "5.60"), ("合计", "—", "约 21.60 Credits")], [170, 150, 190])
    h(doc, "支持的模型", 2)
    add_table(doc, ["品牌", "模型 ID（Model ID）", "模型能力"], models, [90, 205, 235])
    h(doc, "订阅规则", 2)
    h(doc, "购买规则", 3)
    for x in ["支持按月或按年购买，订阅在支付成功后立即生效。", "Token Plan 暂不支持使用代金券购买。", "套餐 Credit 按订阅周期发放；周期结束后，未使用的 Credit 自动清零，不结转至下一周期。"]: item(doc, x)
    h(doc, "有效期", 3)
    para(doc, "套餐有效期以自然月为单位，自购买成功当日开始计算。示例：非闰年 1 月 31 日购买 1 个月套餐，于 3 月 1 日 00:00 到期。")
    h(doc, "续费规则", 3)
    for x in ["支持手动续费和自动续费。月付订阅可续费 1 个月、3 个月或 6 个月，年付订阅每次可续费 1 年；支持多次续费，每次均从原到期日顺延，续费后的最终到期日不得超过当前日期之后 24 个月。", "开启自动续费后，系统将在到期前 7 天尝试扣款；关闭自动续费不影响当前周期使用，到期未续费后订阅失效。"]: item(doc, x)
    h(doc, "升级规则" if personal else "席位变更规则", 3)
    if not personal: item(doc, "加购席位后立即生效，到期时间与当前企业订阅一致；费用和本周期 Credit 按剩余时间折算。")
    item(doc, "升级按当前订阅剩余周期补缴档位差价，支付成功后立即生效且到期时间不变；额度更新为目标档位对应额度，已使用额度保留，当前周期内不支持降级。")
    h(doc, "退款规则", 3)
    para(doc, "新购、续费、" + ("升级" if personal else "加购及升级") + "订单支付成功后不支持退款；关闭自动续费不影响当前周期使用。购买前请确认套餐档位、订阅周期和" + ("购买信息。" if personal else "席位数量。"))
    h(doc, "使用须知", 2)
    for x in ["套餐专属 API Key 仅限本人使用，不得共享或公开泄露。", "套餐额度用尽后，订阅专用 API Key 将停止服务；如需继续调用，可切换至按量付费 API Key。", "套餐支持的模型可能随产品升级调整，实际范围以控制台展示为准。"]: item(doc, x)
    if personal: para(doc, "个人版购买后由本人使用。如需为多名成员统一购买和分配席位，请使用 Token Plan 企业版。")

overview(doc, False)
quick_start(doc, False)
team_management(doc)
faqs = read_faqs()
faq_page(doc, "企业版", faqs)
doc.add_page_break()
overview(doc, True)
quick_start(doc, True)
faq_page(doc, "个人版", faqs)

doc.core_properties.title = "Token Plan 产品文档（完整版）"
doc.core_properties.subject = "Token Plan 页面全部内容：概览、快速开始、团队管理与常见问题"
doc.save(OUT)
print(OUT)
