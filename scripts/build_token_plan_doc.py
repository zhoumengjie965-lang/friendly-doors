from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "Token Plan 产品文档.docx"
BLUE = "2563EB"
NAVY = "17365D"
LIGHT = "EAF2FF"
GRID = "D9E2F3"
MUTED = RGBColor(92, 107, 128)

def font(run, size=10.5, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color if isinstance(color, RGBColor) else RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = tc.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tcMar.find(qn(f"w:{tag}")) or OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa"); tcMar.append(el)
    tc.append(tcMar)

def set_repeat(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)

def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, (h, w) in enumerate(zip(headers, widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); shade(c, BLUE); margins(c)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h), 9.5, True, "FFFFFF")
    set_repeat(t.rows[0])
    for ri, row in enumerate(rows):
        cells=t.add_row().cells
        for i,(text,w) in enumerate(zip(row,widths)):
            cells[i].width=Inches(w); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cells[i])
            if ri%2: shade(cells[i], "F7F9FC")
            p=cells[i].paragraphs[0]; font(p.add_run(str(text)), 9.2, False, "24364B")
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def bullet(doc, text, style="List Bullet"):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(3); font(p.add_run(text),10.5)
    return p

def heading(doc, text, level=1):
    p=doc.add_heading(text, level=level); p.paragraph_format.keep_with_next=True
    return p

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=Inches(.72); sec.left_margin=sec.right_margin=Inches(.78)

styles=doc.styles
normal=styles["Normal"]; normal.font.name="Microsoft YaHei"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); normal.font.size=Pt(10.5)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.22
for name,size,color,before,after in (("Title",30,NAVY,0,8),("Subtitle",13,"536B86",0,18),("Heading 1",17,NAVY,16,8),("Heading 2",13,BLUE,11,5),("Heading 3",11,NAVY,8,4)):
    st=styles[name]; st.font.name="Microsoft YaHei"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=name!="Subtitle"; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)

# Running footer
for s in doc.sections:
    p=s.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    font(p.add_run("Token Plan 产品文档  |  "),8.5,color=MUTED)
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(90); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("TOKEN PLAN"),14,True,BLUE)
p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("产品文档")
p=doc.add_paragraph(style="Subtitle"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("订阅、Credit 计费、支持模型与使用指南")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28)
font(p.add_run("个人版 · 企业版"),11,True,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("更新日期：2026 年 8 月"),9.5,color=MUTED)
doc.add_page_break()

heading(doc,"1. 产品概述")
doc.add_paragraph("Token Plan 是面向个人开发者与企业团队的订阅制 AI 模型服务套餐。用户按月或按年购买套餐，通过专属 API Key 调用套餐支持的模型，并按实际调用量消耗 Credits。企业版按席位购买和分配额度，支持成员、席位及专属 Key 管理；个人版供本人使用。")
heading(doc,"产品特点",2)
for x in ["主流模型一站式使用：统一接入多家模型服务。","统一 Credit 计量：不同模型调用统一折算为 Credits。","套餐额度按周期发放：周期结束后未使用额度自动清零，不结转至下一周期。","企业席位管理：企业管理员可分配、回收席位并查看使用情况。","灵活变更：支持续费、加购席位和升级席位。"]: bullet(doc,x)

heading(doc,"2. 套餐与定价")
doc.add_paragraph("企业版提供 Lite 轻量版、Standard 标准版和 Pro 尊享版三个席位档位，分别适用于日常轻量使用、团队稳定使用和高频调用场景。具体价格与每月额度以购买页面展示为准。")
table(doc,["档位","额度口径","适用场景"],[("Lite 轻量版","按席位、按周期发放","日常轻量使用与个人效率辅助"),("Standard 标准版","按席位、按周期发放","团队日常开发、办公与稳定调用"),("Pro 尊享版","按席位、按周期发放","高频调用与核心生产任务")],[1.45,2.0,3.35])

heading(doc,"3. Credit 抵扣说明")
doc.add_paragraph("模型调用统一折算为 Credit，并从当前订阅或席位的周期额度中扣减。单次消耗由模型类型、Token 用量、思考模式及工具调用等因素动态决定，实际消耗以控制台用量明细为准。")
heading(doc,"示例：qwen3.8-max 单次请求",2)
table(doc,["Token 类型","数量","消耗 Credits"],[("输入 Tokens","8,349","1.67"),("缓存 Tokens","40,794","0.82"),("输出 Tokens","573","0.69"),("合计","—","约 3.18")],[2.3,2.0,2.5])

heading(doc,"4. 支持的模型")
models=[
("DeepSeek","deepseek-v4-pro","推理模型、文本生成"),("DeepSeek","deepseek-v4-flash","推理模型、文本生成"),
("Kimi","kimi-k2.7-code","推理模型、视觉理解、文本生成"),("Kimi","kimi-k2.6","推理模型、视觉理解、文本生成"),("Kimi","kimi-k2.5","推理模型、视觉理解、文本生成"),
("智谱","glm-5.2","推理模型、文本生成"),("智谱","glm-5.1","推理模型、文本生成"),("智谱","glm-5","推理模型、文本生成"),
("阿里","qwen3.8-max-preview","推理模型、视觉理解、文本生成"),("阿里","qwen3.7-max","推理模型、文本生成"),("阿里","qwen3.7-plus","推理模型、视觉理解、文本生成"),("阿里","qwen3.6-plus","推理模型、视觉理解、文本生成"),("阿里","qwen3.6-flash","推理模型、视觉理解、文本生成"),
("阿里","qwen-image-2.0","图片生成"),("阿里","qwen-image-2.0-pro","图片生成"),("阿里","wan2.7-image","图片生成"),("阿里","wan2.7-image-pro","图片生成"),("阿里","happyhorse-1.1-i2v","视频生成"),("阿里","happyhorse-1.1-t2v","视频生成"),("阿里","happyhorse-1.1-r2v","视频生成"),
("MiniMax","MiniMax-M2.5","推理模型、文本生成")]
table(doc,["品牌","模型 ID（Model ID）","模型能力"],models,[1.1,2.35,3.45])

heading(doc,"5. 订阅规则")
heading(doc,"购买规则",2)
for x in ["支持按月或按年购买，订阅在支付成功后立即生效。","套餐 Credit 按订阅周期发放；周期结束后，未使用的 Credit 自动清零，不结转至下一周期。","Token Plan 暂不支持使用代金券购买。"]: bullet(doc,x)
heading(doc,"续费规则",2)
for x in ["支持手动续费和自动续费；月付订阅可选择续费 1 个月、3 个月或 6 个月，年付订阅每次可续费 1 年；订阅支持多次续费，每次均从原到期日顺延，续费后的最终到期日不得超过当前日期之后 24 个月。","开启自动续费后，系统将在到期前 7 天尝试扣款；关闭自动续费不影响当前周期使用，到期未续费后订阅失效。"]: bullet(doc,x)
heading(doc,"席位变更规则",2)
for x in ["加购席位支付成功后立即生效，到期时间与当前企业订阅保持一致；费用和本周期 Credit 按剩余天数折算，下周期恢复完整周期额度。","升级按当前订阅剩余周期补缴档位差价，支付成功后立即生效，到期时间不变。","升级后额度更新为目标档位对应额度，已使用额度保留；当前周期内不支持降级。"]: bullet(doc,x)
heading(doc,"退款规则",2)
doc.add_paragraph("Token Plan 企业版新购、续费、加购及升级成功后均不支持退款。")

heading(doc,"6. 快速开始")
heading(doc,"企业版",2)
for x in ["企业管理员完成套餐购买。","在企业订阅中为成员分配席位。","成员创建或获取席位专属 API Key。","使用专属 API Key 调用支持的模型。","在订阅与用量页面查看剩余额度和消耗明细。"]: bullet(doc,x,"List Number")
heading(doc,"个人版",2)
for x in ["选择套餐与订阅周期并完成支付。","创建 Token Plan 专属 API Key。","调用支持的模型并查看 Credit 消耗。"]: bullet(doc,x,"List Number")

heading(doc,"7. 企业团队管理")
for x in ["席位分配：企业管理员可将空闲席位分配给企业成员。","席位回收：回收后成员不再拥有该席位的使用权限，相关专属 Key 停止使用。","专属 Key：每个席位仅绑定一名成员，不支持多人共用。","使用查询：管理员可查看企业订阅、席位状态及额度使用情况。"]: bullet(doc,x)

heading(doc,"8. 常见问题")
faqs=[("套餐额度用尽后还能继续调用吗？","不能。套餐额度用尽后，该 Token Plan Key 将停止服务；如需继续使用，可切换至按量付费 Key。"),("未使用的 Credit 会结转吗？","不会。Credit 在每个订阅周期结束后自动清零。"),("可以取消自动续费吗？","可以。关闭自动续费不影响当前周期使用，订阅将在到期且未续费后失效。"),("支持退款或降级吗？","购买成功后不支持退款；当前订阅周期内不支持席位降级。"),("加购席位如何计费？","按企业当前订阅的剩余天数计价，本周期额度同步按剩余天数折算，并与企业订阅同时到期。")]
for q,a in faqs:
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2); font(p.add_run(q),10.5,True,NAVY)
    doc.add_paragraph(a)

doc.core_properties.title="Token Plan 产品文档"; doc.core_properties.subject="Token Plan 产品说明、支持模型、订阅规则与使用指南"
doc.save(OUT)
print(OUT)
